import docx
from datetime import datetime

data = {
    'date': datetime.utcnow(),
    'num_client': 5,
    'total': 376,
}

doc = docx.Document()
doc.add_heading('Raport z dnia pracy',0)
par = doc.add_paragraph('Data: ')
par.add_run(str(data['date'])).italic = True
par = doc.add_paragraph('Liczba minut w pracy: ')
par.add_run(str(data['total'])).bold = True

doc.save('raport-word.docx')